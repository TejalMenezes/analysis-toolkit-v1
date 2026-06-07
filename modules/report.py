
"""Report engine for Smart Analysis Reporter.

Holds the report state (cover + items), renders themed matplotlib charts to PNG
for embedding, builds inference sentences, and exports a self-contained HTML
report and an editable DOCX document.
"""

import io
import base64
import uuid
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
import streamlit as st

from modules.ui import (
    ORANGE, ORANGE_DARK, ORANGE_SOFT, ORANGE_LINE, INK, MUTED, PALETTE,
    APP_NAME, AUTHOR_NAME, AUTHOR_ID,
)


# ───────────────────────── state ─────────────────────────

def _default_cover():
    return {
        "title": "Smart Analysis Report",
        "subtitle": "Student Performance Prediction",
        "author": AUTHOR_NAME,
        "author_id": AUTHOR_ID,
        "date": date.today().isoformat(),
        "summary": (
            "This report summarises an exploratory and inferential analysis of the "
            "dataset. Each section pairs a chart or table with a short, editable "
            "interpretation — feel free to refine the wording before exporting."
        ),
    }


def init_report():
    if "report" not in st.session_state:
        st.session_state["report"] = {"cover": _default_cover(), "items": []}


def get_report():
    init_report()
    return st.session_state["report"]


def reset_report():
    st.session_state["report"] = {"cover": _default_cover(), "items": []}


def add_item(title, inference="", image=None, table=None, kind="block"):
    """Append an item. image=PNG bytes, table=DataFrame. Returns the item id."""
    rep = get_report()
    item = {
        "id": uuid.uuid4().hex[:8],
        "kind": kind,
        "title": title,
        "inference": inference,
        "image": image,
        "table": table,
    }
    rep["items"].append(item)
    return item["id"]


def already_added(title):
    return any(it["title"] == title for it in get_report()["items"])


def remove_item(item_id):
    rep = get_report()
    rep["items"] = [it for it in rep["items"] if it["id"] != item_id]


def move_item(item_id, delta):
    rep = get_report()
    items = rep["items"]
    idx = next((i for i, it in enumerate(items) if it["id"] == item_id), None)
    if idx is None:
        return
    new = idx + delta
    if 0 <= new < len(items):
        items[idx], items[new] = items[new], items[idx]


def add_to_report_button(title, *, image=None, table=None, inference="", key=None):
    """A small UI control used on analysis pages to capture the current
    chart/table into the report."""
    if already_added(title):
        st.caption(f"✓ \"{title}\" is already in your report.")
        return
    if st.button("➕ Add to report", key=key or f"add_{title}"):
        add_item(title, inference=inference, image=image, table=table)
        st.toast(f"Added \"{title}\" to report", icon="📝")
        st.rerun()


# ───────────────────── matplotlib theme ─────────────────────

def _style_ax(ax, title=None, xlabel=None, ylabel=None):
    ax.set_facecolor("#FFFFFF")
    for s in ("top", "right"):
        ax.spines[s].set_visible(False)
    for s in ("left", "bottom"):
        ax.spines[s].set_color("#D9D9D9")
    ax.tick_params(colors=MUTED, labelsize=8)
    ax.grid(True, color="#F0F0F0", linewidth=0.8)
    ax.set_axisbelow(True)
    if title:
        ax.set_title(title, color=INK, fontsize=11, fontweight="bold", pad=10)
    if xlabel:
        ax.set_xlabel(xlabel, color=MUTED, fontsize=9)
    if ylabel:
        ax.set_ylabel(ylabel, color=MUTED, fontsize=9)


def _png(fig):
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=130, bbox_inches="tight", facecolor="white")
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def hist_png(series, title=""):
    s = pd.to_numeric(pd.Series(series), errors="coerce").dropna()
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    nb = max(5, int(np.sqrt(len(s)))) if len(s) else 5
    ax.hist(s, bins=nb, color=ORANGE, edgecolor="white", alpha=0.9)
    _style_ax(ax, title, ylabel="Count")
    return _png(fig)


def box_png(series, title=""):
    s = pd.to_numeric(pd.Series(series), errors="coerce").dropna()
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    bp = ax.boxplot(s, vert=False, patch_artist=True, widths=0.5)
    for b in bp["boxes"]:
        b.set(facecolor=ORANGE_SOFT, edgecolor=ORANGE, linewidth=1.6)
    for w in bp["whiskers"] + bp["caps"]:
        w.set(color=ORANGE_DARK, linewidth=1.2)
    for m in bp["medians"]:
        m.set(color=ORANGE_DARK, linewidth=2)
    for fl in bp["fliers"]:
        fl.set(marker="o", markerfacecolor=ORANGE, markersize=4, alpha=0.5, markeredgecolor="none")
    _style_ax(ax, title)
    ax.set_yticks([])
    return _png(fig)


def bar_png(labels, values, title="", xlabel="", ylabel="Count"):
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    cols = [PALETTE[i % len(PALETTE)] for i in range(len(labels))]
    ax.bar([str(l) for l in labels], values, color=cols, edgecolor="white")
    _style_ax(ax, title, xlabel=xlabel, ylabel=ylabel)
    if len(labels) > 6:
        plt.setp(ax.get_xticklabels(), rotation=35, ha="right")
    return _png(fig)


def scatter_png(x, y, xlabel="", ylabel="", title=""):
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    ax.scatter(x, y, color=ORANGE, alpha=0.45, s=18, edgecolors="none")
    _style_ax(ax, title, xlabel=xlabel, ylabel=ylabel)
    return _png(fig)


def regression_png(x, y, x_line, y_line, xlabel="", ylabel="", title="Regression fit"):
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    ax.scatter(x, y, color=ORANGE, alpha=0.35, s=16, edgecolors="none")
    ax.plot(x_line, y_line, color=ORANGE_DARK, linewidth=2.4)
    _style_ax(ax, title, xlabel=xlabel, ylabel=ylabel)
    return _png(fig)


def heatmap_png(corr_df, title="Correlation matrix"):
    fig, ax = plt.subplots(figsize=(5.6, 4.4))
    data = corr_df.values.astype(float)
    cmap = matplotlib.colors.LinearSegmentedColormap.from_list(
        "sar", ["#3B6FB6", "#FFFFFF", ORANGE_DARK]
    )
    im = ax.imshow(data, cmap=cmap, vmin=-1, vmax=1, aspect="auto")
    ax.set_xticks(range(len(corr_df.columns)))
    ax.set_yticks(range(len(corr_df.index)))
    ax.set_xticklabels(corr_df.columns, rotation=40, ha="right", fontsize=7, color=MUTED)
    ax.set_yticklabels(corr_df.index, fontsize=7, color=MUTED)
    for i in range(data.shape[0]):
        for j in range(data.shape[1]):
            v = data[i, j]
            ax.text(j, i, f"{v:.2f}", ha="center", va="center",
                    fontsize=6.5, color="#fff" if abs(v) > 0.55 else INK)
    ax.set_title(title, color=INK, fontsize=11, fontweight="bold", pad=10)
    fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
    return _png(fig)


def qq_png(theo, sample, ref_x, ref_y, title="Q-Q plot (normal)"):
    fig, ax = plt.subplots(figsize=(5.4, 3.2))
    ax.scatter(theo, sample, color=ORANGE, alpha=0.5, s=16, edgecolors="none")
    ax.plot(ref_x, ref_y, color=ORANGE_DARK, linewidth=2)
    _style_ax(ax, title, xlabel="Theoretical quantiles", ylabel="Sample quantiles")
    return _png(fig)


def line_png(x, y, xlabel="", ylabel="", title=""):
    fig, ax = plt.subplots(figsize=(5.8, 3.0))
    ax.plot(x, y, color=ORANGE, linewidth=1.8)
    _style_ax(ax, title, xlabel=xlabel, ylabel=ylabel)
    return _png(fig)


# ───────────────────── inference builders ─────────────────────

def describe_inference(name, s):
    s = pd.to_numeric(pd.Series(s), errors="coerce").dropna()
    if len(s) == 0:
        return f"No numeric data available for {name}."
    skew = s.skew()
    shape = ("roughly symmetric" if abs(skew) < 0.5
             else "right-skewed" if skew > 0 else "left-skewed")
    return (
        f"{name} averages {s.mean():.2f} (median {s.median():.2f}, SD {s.std():.2f}), "
        f"ranging from {s.min():.2f} to {s.max():.2f}. The distribution is {shape} "
        f"(skewness {skew:.2f})."
    )


def correlation_inference(corr_df):
    pairs = []
    cols = corr_df.columns
    for i in range(len(cols)):
        for j in range(i + 1, len(cols)):
            pairs.append((cols[i], cols[j], corr_df.iloc[i, j]))
    if not pairs:
        return "Not enough numeric variables to assess correlation."
    pairs.sort(key=lambda p: abs(p[2]), reverse=True)
    a, b, r = pairs[0]
    strength = ("a strong" if abs(r) >= 0.7 else "a moderate" if abs(r) >= 0.4
                else "a weak")
    direction = "positive" if r > 0 else "negative"
    return (
        f"The strongest linear relationship is between {a} and {b} "
        f"(r = {r:.2f}), {strength} {direction} association. "
        "Variables with high correlation may carry overlapping information."
    )


def regression_inference(x_var, y_var, r):
    sig = r["p_slope"] < 0.05
    return (
        f"A simple linear model predicts {y_var} from {x_var} as "
        f"{r['equation']}. It explains {r['r2'] * 100:.1f}% of the variance "
        f"(R² = {r['r2']:.3f}). The slope is "
        f"{'statistically significant' if sig else 'not statistically significant'} "
        f"(p {'< 0.001' if r['p_slope'] < 0.001 else f'= {r['p_slope']:.3f}'}), so each "
        f"one-unit rise in {x_var} is associated with a change of "
        f"{r['slope']:.3f} in {y_var}."
    )


def test_inference(reject, msg_yes, msg_no):
    return ("Reject H₀ — " + msg_yes) if reject else ("Fail to reject H₀ — " + msg_no)


# ───────────────────── exporters ─────────────────────

def _b64(png):
    return base64.b64encode(png).decode("ascii")


def build_html(report):
    cover = report["cover"]
    items = report["items"]

    blocks = []
    for n, it in enumerate(items, 1):
        img = (f'<img src="data:image/png;base64,{_b64(it["image"])}" '
               f'style="max-width:100%;border:1px solid #eee;border-radius:10px;margin:8px 0"/>'
               if it.get("image") else "")
        tbl = ""
        if it.get("table") is not None:
            tbl = it["table"].to_html(
                classes="sar-table", index=True, border=0, float_format=lambda v: f"{v:.3f}"
            )
        inf = (f'<p class="sar-inf">{it["inference"]}</p>' if it.get("inference") else "")
        blocks.append(
            f'<section class="sar-block"><h2><span class="num">{n}</span>{it["title"]}</h2>'
            f'{img}{tbl}{inf}</section>'
        )

    body = "\n".join(blocks)

    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>{cover['title']}</title>
<style>
  body {{ font-family:'Segoe UI',system-ui,sans-serif; color:{INK}; margin:0; background:#fff; }}
  .wrap {{ max-width:880px; margin:0 auto; padding:0 38px 60px; }}
  .cover {{ background:linear-gradient(135deg,{ORANGE},#FF9E40); color:#fff;
            padding:64px 48px 48px; border-radius:0 0 26px 26px; margin-bottom:34px; }}
  .cover .app {{ font-size:13px; letter-spacing:3px; text-transform:uppercase; opacity:.9; }}
  .cover h1 {{ font-size:40px; margin:14px 0 6px; font-weight:800; }}
  .cover .sub {{ font-size:18px; opacity:.95; }}
  .cover .meta {{ margin-top:26px; font-size:14px; }}
  .cover .meta b {{ font-weight:700; }}
  .summary {{ background:{ORANGE_SOFT}; border-left:5px solid {ORANGE};
              padding:18px 22px; border-radius:12px; margin:0 38px 30px; font-size:14.5px; line-height:1.6; }}
  .sar-block {{ margin:0 38px 30px; padding-bottom:24px; border-bottom:1px solid #eee; }}
  .sar-block h2 {{ font-size:19px; color:{ORANGE_DARK}; display:flex; align-items:center; gap:12px; }}
  .sar-block h2 .num {{ background:{ORANGE}; color:#fff; width:28px; height:28px; border-radius:8px;
                        display:inline-flex; align-items:center; justify-content:center; font-size:14px; }}
  .sar-inf {{ font-size:14px; line-height:1.65; color:#333; margin-top:10px;
              background:#FAFAFA; padding:12px 16px; border-radius:10px; }}
  table.sar-table {{ border-collapse:collapse; width:100%; font-size:12.5px; margin:8px 0; }}
  table.sar-table th {{ background:{ORANGE_SOFT}; color:{ORANGE_DARK}; text-align:left;
                        padding:7px 10px; border-bottom:2px solid {ORANGE_LINE}; }}
  table.sar-table td {{ padding:6px 10px; border-bottom:1px solid #eee; }}
  .foot {{ text-align:center; color:{MUTED}; font-size:12px; margin-top:40px; }}
  @media print {{ .cover {{ -webkit-print-color-adjust:exact; print-color-adjust:exact; }} }}
</style></head><body>
  <div class="cover">
    <div class="app">{APP_NAME}</div>
    <h1>{cover['title']}</h1>
    <div class="sub">{cover['subtitle']}</div>
    <div class="meta">Prepared by <b>{cover['author']}</b> &nbsp;·&nbsp; ID {cover['author_id']}
      &nbsp;·&nbsp; {cover['date']}</div>
  </div>
  <div class="summary"><b>Executive summary.</b> {cover['summary']}</div>
  <div class="wrap" style="padding-top:0">{body}
    <div class="foot">Generated with {APP_NAME} · {cover['author']} ({cover['author_id']})</div>
  </div>
</body></html>"""


def _rl_styles():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT
    from reportlab.lib.colors import HexColor

    ss = getSampleStyleSheet()
    styles = {
        "h2": ParagraphStyle("sar_h2", parent=ss["Heading2"], textColor=HexColor(ORANGE_DARK),
                             fontSize=15, spaceBefore=14, spaceAfter=6),
        "body": ParagraphStyle("sar_body", parent=ss["BodyText"], fontSize=10.5,
                               leading=15, textColor=HexColor("#333333")),
        "inf": ParagraphStyle("sar_inf", parent=ss["BodyText"], fontSize=10.5, leading=15,
                              textColor=HexColor("#333333"), backColor=HexColor("#FAFAFA"),
                              borderPadding=8, leftIndent=2, spaceBefore=6),
        "cover_app": ParagraphStyle("sar_capp", parent=ss["BodyText"], textColor=HexColor("#FFFFFF"),
                                    fontSize=11, leading=14),
        "cover_title": ParagraphStyle("sar_ctitle", parent=ss["Title"], textColor=HexColor("#FFFFFF"),
                                      fontSize=30, leading=34, spaceBefore=8, spaceAfter=4, alignment=TA_LEFT),
        "cover_sub": ParagraphStyle("sar_csub", parent=ss["BodyText"], textColor=HexColor("#FFFFFF"),
                                    fontSize=15, leading=19),
        "cover_meta": ParagraphStyle("sar_cmeta", parent=ss["BodyText"], textColor=HexColor("#FFFFFF"),
                                     fontSize=11, leading=16, spaceBefore=14),
    }
    return styles


def _rl_image(png, max_w):
    import io as _io
    from reportlab.platypus import Image
    from reportlab.lib.utils import ImageReader
    iw, ih = ImageReader(_io.BytesIO(png)).getSize()
    w = min(max_w, iw)
    return Image(_io.BytesIO(png), width=w, height=w * ih / iw)


def _rl_table(df, max_w):
    from reportlab.platypus import Table, TableStyle
    from reportlab.lib.colors import HexColor

    header = [str(df.index.name or "")] + [str(c) for c in df.columns]
    rows = [header]
    for idx, row in df.iterrows():
        cells = [str(idx)]
        for c in df.columns:
            v = row[c]
            cells.append(f"{v:.3f}" if isinstance(v, (int, float, np.floating)) and not isinstance(v, bool)
                         else str(v))
        rows.append(cells)

    ncol = len(header)
    col_w = max_w / ncol
    t = Table(rows, colWidths=[col_w] * ncol, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(ORANGE_SOFT)),
        ("TEXTCOLOR", (0, 0), (-1, 0), HexColor(ORANGE_DARK)),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("LINEBELOW", (0, 0), (-1, 0), 1, HexColor(ORANGE_LINE)),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [HexColor("#FFFFFF"), HexColor("#FBFBFB")]),
        ("LINEBELOW", (0, 1), (-1, -1), 0.4, HexColor("#EEEEEE")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def build_pdf(report):
    """A polished, self-contained PDF of the report (orange cover + sections)."""
    import io as _io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import cm
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, KeepTogether)

    cover = report["cover"]
    items = report["items"]
    S = _rl_styles()

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=1.4 * cm, bottomMargin=1.4 * cm,
                            leftMargin=1.6 * cm, rightMargin=1.6 * cm,
                            title=cover["title"], author=cover["author"])
    avail_w = doc.width
    story = []

    # cover band
    cover_inner = [
        [Paragraph(APP_NAME.upper(), S["cover_app"])],
        [Paragraph(cover["title"], S["cover_title"])],
        [Paragraph(cover["subtitle"], S["cover_sub"])],
        [Paragraph(f"Prepared by <b>{cover['author']}</b> &nbsp;·&nbsp; ID {cover['author_id']} "
                   f"&nbsp;·&nbsp; {cover['date']}", S["cover_meta"])],
    ]
    band = Table(cover_inner, colWidths=[avail_w])
    band.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor(ORANGE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 22),
        ("RIGHTPADDING", (0, 0), (-1, -1), 22),
        ("TOPPADDING", (0, 0), (0, 0), 26),
        ("BOTTOMPADDING", (-1, -1), (-1, -1), 24),
    ]))
    story += [band, Spacer(1, 14)]

    # executive summary
    summ = Table([[Paragraph("<b>Executive summary.</b> " + cover["summary"], S["body"])]],
                 colWidths=[avail_w])
    summ.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor(ORANGE_SOFT)),
        ("LINEBEFORE", (0, 0), (0, -1), 3, HexColor(ORANGE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 14),
        ("RIGHTPADDING", (0, 0), (-1, -1), 14),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story += [summ, Spacer(1, 16)]

    for n, it in enumerate(items, 1):
        block = [Paragraph(f"{n}. {it['title']}", S["h2"])]
        if it.get("image"):
            block += [_rl_image(it["image"], avail_w), Spacer(1, 6)]
        if it.get("table") is not None:
            block += [_rl_table(it["table"], avail_w), Spacer(1, 6)]
        if it.get("inference"):
            block += [Paragraph(it["inference"], S["inf"])]
        block += [Spacer(1, 14)]
        story.append(KeepTogether(block) if not it.get("image") else block[0])
        if it.get("image"):
            story += block[1:]

    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def build_docx(report):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    cover = report["cover"]
    doc = Document()

    app_p = doc.add_paragraph(APP_NAME.upper())
    app_p.runs[0].font.size = Pt(11)
    app_p.runs[0].font.color.rgb = RGBColor(0xF5, 0x7C, 0x00)

    h = doc.add_heading(cover["title"], level=0)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    sub = doc.add_paragraph(cover["subtitle"])
    sub.runs[0].font.size = Pt(14)
    doc.add_paragraph(
        f"Prepared by {cover['author']}  ·  ID {cover['author_id']}  ·  {cover['date']}"
    )

    doc.add_heading("Executive summary", level=2)
    doc.add_paragraph(cover["summary"])

    for n, it in enumerate(report["items"], 1):
        doc.add_heading(f"{n}. {it['title']}", level=2)
        if it.get("image"):
            doc.add_picture(io.BytesIO(it["image"]), width=Inches(5.6))
        if it.get("table") is not None:
            t = it["table"]
            tbl = doc.add_table(rows=1, cols=len(t.columns) + 1)
            tbl.style = "Light Grid Accent 6"
            hdr = tbl.rows[0].cells
            hdr[0].text = str(t.index.name or "")
            for j, c in enumerate(t.columns):
                hdr[j + 1].text = str(c)
            for idx, row in t.iterrows():
                cells = tbl.add_row().cells
                cells[0].text = str(idx)
                for j, c in enumerate(t.columns):
                    val = row[c]
                    cells[j + 1].text = f"{val:.3f}" if isinstance(val, (int, float, np.floating)) else str(val)
        if it.get("inference"):
            doc.add_paragraph(it["inference"])

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
