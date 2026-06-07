
"""Professional documentation generators (PDF + Word).

Each document is defined once as a content *model* — a list of sections, each a
title plus typed blocks (heading, paragraph, bullets, problem-statement box,
figure, caption, table). The same model is rendered to:

* PDF  (ReportLab)      — cover page, table of contents, page headers/footers
                          with page numbers, each section on its own page.
* DOCX (python-docx)    — matching cover, Word TOC field, heading styles, page
                          breaks, footer page numbers.

Two documents are produced from the dataset:
  * Analysis Documentation — a guided statistical study.
  * System Documentation   — design, tech stack, dataset, analysis, conclusions.
"""

import io
import re
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
import numpy as np
import pandas as pd

from modules import report as R
from modules.ui import (ORANGE, ORANGE_DARK, ORANGE_SOFT, ORANGE_LINE, INK, MUTED,
                        APP_NAME, AUTHOR_NAME, AUTHOR_ID)
from modules.descriptive import describe_all
from modules.frequency import categorical_frequency
from modules.normality import normality_report, qq_data
from modules.regression import simple_linear_regression
from modules.tests import independent_ttest, effect_label_d


# ───────────────── content model (block constructors) ─────────────────

def H2(text):       return ("h2", text)
def P(text):        return ("para", text)
def BUL(items):     return ("bullets", items)
def PROB(text):     return ("problem", text)
def FIG(png, frac=0.82): return ("figure", png, frac)
def CAP(text):      return ("caption", text)
def TBL(df):        return ("table", df)
def CODE(text):     return ("code", text.strip("\n"))


def SEC(title, *blocks):
    return {"title": title, "blocks": [b for b in blocks if b is not None]}


# ───────────────── dataset structure helpers ─────────────────

def _numeric(df):
    return list(df.select_dtypes(include="number").columns)


def _categorical(df):
    return list(df.select_dtypes(include=["object", "category", "bool"]).columns)


def _pick_target(df, numeric):
    for cand in ("exam_score", "score", "target", "result"):
        for c in numeric:
            if c.lower() == cand:
                return c
    return max(numeric, key=lambda c: df[c].nunique()) if numeric else None


# ───────────────── architecture diagram ─────────────────

def architecture_png():
    fig, ax = plt.subplots(figsize=(7.6, 3.4))
    ax.axis("off"); ax.set_xlim(0, 10); ax.set_ylim(0, 4)

    def box(x, y, w, h, label, fill):
        ax.add_patch(mpatches.FancyBboxPatch(
            (x, y), w, h, boxstyle="round,pad=0.04,rounding_size=0.12",
            linewidth=1.4, edgecolor=ORANGE_DARK, facecolor=fill))
        ax.text(x + w / 2, y + h / 2, label, ha="center", va="center",
                fontsize=8.5, color=INK, weight="bold")

    box(0.2, 1.4, 1.7, 1.2, "Dataset\n(CSV cache)", "#FFFFFF")
    box(2.4, 1.4, 1.9, 1.2, "Analysis modules\n(stats / tests /\nregression)", ORANGE_SOFT)
    box(4.8, 2.3, 2.0, 1.0, "Streamlit pages\n(interactive UI)", "#FFFFFF")
    box(4.8, 0.5, 2.0, 1.0, "Auto-report\nbuilder", ORANGE_SOFT)
    box(7.3, 1.4, 2.4, 1.2, "Report engine →\nPDF · DOCX · HTML", ORANGE)
    ax.text(8.5, 2.05, "exports", ha="center", fontsize=7, color="#fff")

    def arrow(x1, y1, x2, y2):
        ax.annotate("", xy=(x2, y2), xytext=(x1, y1),
                    arrowprops=dict(arrowstyle="-|>", color=MUTED, lw=1.3))

    arrow(1.9, 2.0, 2.4, 2.0); arrow(4.3, 2.2, 4.8, 2.6); arrow(4.3, 1.8, 4.8, 1.0)
    arrow(6.8, 2.7, 7.4, 2.2); arrow(6.8, 1.0, 7.4, 1.7)
    ax.set_title("System architecture & data flow", color=INK, fontsize=11, weight="bold", pad=8)
    return R._png(fig)


def _key_findings(df):
    numeric, categorical = _numeric(df), _categorical(df)
    target = _pick_target(df, numeric)
    out = []
    if len(numeric) >= 2:
        corr = df[numeric].corr()
        pairs = [(corr.columns[i], corr.columns[j], corr.iloc[i, j])
                 for i in range(len(corr.columns)) for j in range(i + 1, len(corr.columns))]
        a, b, rv = max(pairs, key=lambda p: abs(p[2]))
        out.append(f"The strongest linear relationship is between <b>{a}</b> and <b>{b}</b> (r = {rv:.2f}).")
        if target:
            ct = corr[target].drop(target).abs()
            if len(ct):
                p = ct.idxmax()
                r = simple_linear_regression(df, p, target)
                if r:
                    out.append(f"A simple model of <b>{target}</b> on <b>{p}</b> explains "
                               f"{r['r2']*100:.1f}% of its variance (R² = {r['r2']:.2f}).")
    if categorical and target:
        cat = categorical[0]
        g = list(df[cat].dropna().unique())
        if len(g) == 2:
            r = independent_ttest(df[df[cat] == g[0]][target], df[df[cat] == g[1]][target])
            verdict = "a statistically significant" if r["p"] < 0.05 else "no significant"
            out.append(f"There is {verdict} difference in <b>{target}</b> between <b>{cat}</b> "
                       f"groups (p {'< 0.001' if r['p'] < 0.001 else f'= {r['p']:.3f}'}).")
    if target:
        nr = normality_report(df[target])
        if nr["W"] is not None:
            out.append(f"<b>{target}</b> is "
                       f"{'approximately normally distributed' if nr['normal'] else 'not normally distributed'} "
                       f"(Shapiro-Wilk W = {nr['W']}).")
    return out


# ═══════════════════ DOCUMENT MODELS ═══════════════════

def build_analysis_model(df, name):
    numeric, categorical = _numeric(df), _categorical(df)
    target = _pick_target(df, numeric)
    sections = []

    sections.append(SEC(
        "1. Introduction",
        H2("1.1  Purpose of this document"),
        P("This document presents a structured statistical analysis of the "
          f"<b>{name}</b> dataset. It is written to be read top-to-bottom: each section poses "
          "a concrete question, states the method used to answer it, presents the evidence (a "
          "chart or table), and closes with an interpretation. The goal is a defensible, "
          "reproducible understanding of what drives student outcomes."),
        H2("1.2  Problem statement"),
        PROB("Educational institutions need to understand which study-related factors are "
             "associated with student exam performance, whether outcomes differ across student "
             "groups, and whether performance can be predicted from measurable behaviours. This "
             f"study interrogates the {name} dataset ({len(df):,} students, {df.shape[1]} "
             "variables) to answer those questions."),
        H2("1.3  Objectives"),
        BUL(["Describe the central tendency, spread and shape of the key variables.",
             "Test whether exam performance differs significantly across categorical groups.",
             "Measure which factors are most strongly correlated with performance.",
             "Build and evaluate a predictive linear-regression model.",
             "Synthesise the findings into an overall understanding of the dataset."]),
        H2("1.4  Methodology & tools"),
        P("Analyses were produced with the Smart Analysis Reporter toolkit (Python, pandas, "
          "SciPy, statsmodels). Descriptive statistics, frequency analysis, correlation, "
          "independent-samples t-tests and ordinary-least-squares regression were applied, each "
          "with a significance level of α = 0.05."),
    ))

    desc = describe_all(df, numeric)
    sections.append(SEC(
        "2. Descriptive Analysis",
        PROB("What are the typical values, variability and distribution shape of the study "
             "habits and performance measures in the dataset?"),
        H2("2.1  Method"),
        P("Summary statistics (mean, median, mode, variance, standard deviation, quartiles, "
          "skewness and kurtosis) were computed for every metric variable."),
        H2("2.2  Results"),
        TBL(desc),
        FIG(R.hist_png(df[target], title=f"Distribution of {target}")) if target else None,
        CAP(f"Figure 2.1 — Distribution of the target variable, {target}.") if target else None,
        H2("2.3  Interpretation"),
        P(R.describe_inference(target, df[target]) if target else
          "The table above establishes the baseline profile of every metric variable."),
    ))

    cat_blocks = [PROB("Do students differ in exam performance across categorical groups — "
                       "for example, between placement outcomes?")]
    if categorical:
        cat = categorical[0]
        freq = categorical_frequency(df[cat])
        top = freq.iloc[0]
        cat_blocks += [
            H2("3.1  Group composition"),
            FIG(R.bar_png(freq["Value"], freq["Abs. Frequency"],
                          title=f"{cat} distribution", xlabel=cat)),
            CAP(f"Figure 3.1 — Distribution of {cat}."),
            P(f"{cat} has {len(freq)} categories; the most common is \"{top['Value']}\" "
              f"({top['Rel. Frequency %']}% of records)."),
        ]
        groups = list(df[cat].dropna().unique())
        if target and len(groups) == 2:
            g1, g2 = df[df[cat] == groups[0]][target], df[df[cat] == groups[1]][target]
            r = independent_ttest(g1, g2)
            sig = r["p"] < 0.05
            p_str = "< 0.001" if r["p"] < 0.001 else f"= {r['p']:.3f}"
            comp = pd.DataFrame({"Mean": [r["mean1"], r["mean2"]], "SD": [r["sd1"], r["sd2"]],
                                 "N": [r["n1"], r["n2"]]}, index=[str(groups[0]), str(groups[1])])
            comp.index.name = str(cat)
            cat_blocks += [
                H2("3.2  Hypothesis test (independent-samples t-test)"),
                P(f"H₀: the mean {target} is equal across {cat} groups. "
                  f"H₁: the means differ."),
                TBL(comp),
                H2("3.3  Interpretation"),
                P(R.test_inference(
                    sig,
                    f"{target} differs significantly between \"{groups[0]}\" and \"{groups[1]}\" "
                    f"(t = {r['stat']:.2f}, p {p_str}, Cohen's d = {r['cohen_d']:.2f}, "
                    f"{effect_label_d(r['cohen_d'])} effect).",
                    f"{target} does not differ significantly between \"{groups[0]}\" and "
                    f"\"{groups[1]}\" (p {p_str}).")),
            ]
    sections.append(SEC("3. Categorical Testing", *cat_blocks))

    predictor = None
    corr_blocks = [PROB("Which variables are most strongly associated with exam performance, "
                        "and could therefore serve as predictors?")]
    if len(numeric) >= 2:
        corr = df[numeric].corr()
        if target:
            ct = corr[target].drop(target).abs()
            predictor = ct.idxmax() if len(ct) else None
        corr_blocks += [
            H2("4.1  Method"),
            P("Pearson correlation coefficients were computed for all metric variable pairs "
              "(range −1 to +1)."),
            H2("4.2  Results"),
            FIG(R.heatmap_png(corr)),
            CAP("Figure 4.1 — Correlation matrix of the metric variables."),
            H2("4.3  Interpretation"),
            P(R.correlation_inference(corr)),
        ]
    sections.append(SEC("4. Correlation & Relationships", *corr_blocks))

    reg_blocks = [PROB("Can exam performance be predicted from its strongest single factor, and "
                       "how much of the variation does that factor explain?")]
    if target and predictor:
        r = simple_linear_regression(df, predictor, target)
        if r:
            xnew = float(df[predictor].mean() + df[predictor].std())
            pred = r["intercept"] + r["slope"] * xnew
            reg_blocks += [
                H2("5.1  Method"),
                P(f"A simple ordinary-least-squares linear regression was fitted, modelling "
                  f"{target} as a function of {predictor}."),
                H2("5.2  Results"),
                FIG(R.regression_png(r["x"], r["y"], r["x_line"], r["y_line"],
                                     xlabel=predictor, ylabel=target,
                                     title=f"{target} vs {predictor}")),
                CAP(f"Figure 5.1 — Regression of {target} on {predictor}."),
                H2("5.3  Interpretation"),
                P(R.regression_inference(predictor, target, r)),
                P(f"As a worked example, a student one standard deviation above the mean "
                  f"{predictor} ({xnew:.1f}) is predicted to score <b>{pred:.1f}</b> on {target}."),
            ]
    sections.append(SEC("5. Predictive Trends — Linear Regression", *reg_blocks))

    sections.append(SEC(
        "6. Synthesis & Conclusions",
        H2("6.1  Bringing the analysis together"),
        P("Each section answered one question, and together they form a coherent picture. The "
          "descriptive profile set expectations; categorical testing showed where groups "
          "genuinely differ; correlation identified the factors that move with performance; and "
          "the regression turned the strongest of those into a predictive rule."),
        H2("6.2  Key findings"),
        BUL(_key_findings(df)),
        H2("6.3  Limitations & next steps"),
        BUL(["A single-predictor model is a baseline — a multiple-regression model using "
             "several factors together would likely explain more variance.",
             "Correlation does not imply causation; a controlled study is needed to confirm drivers.",
             "Findings describe this sample; external validation on new cohorts is recommended."]),
    ))

    meta = {"title": "Analysis Documentation", "doc_kind": "Statistical Analysis Report",
            "subtitle": name}
    return meta, sections


def _feature(title, purpose, inputs, methods, outputs, module):
    """A granular feature entry: heading + structured bullet list."""
    return [
        H2(title),
        BUL([
            f"<b>Purpose.</b> {purpose}",
            f"<b>User inputs / controls.</b> {inputs}",
            f"<b>Methods & computation.</b> {methods}",
            f"<b>Outputs.</b> {outputs}",
            f"<b>Backing module.</b> {module}",
        ]),
    ]


def build_system_model(df, name):
    numeric, categorical = _numeric(df), _categorical(df)
    target = _pick_target(df, numeric)
    predictor = None
    if len(numeric) >= 2 and target:
        ct = df[numeric].corr()[target].drop(target).abs()
        predictor = ct.idxmax() if len(ct) else None
    sections = []

    # 1 ── Introduction & purpose
    sections.append(SEC(
        "1. Introduction & Purpose",
        H2("1.1  About this document"),
        P("This document describes the design and implementation of <b>Smart Analysis "
          "Reporter</b>, a web-based statistical analysis and reporting toolkit. Its focus is "
          "the <b>toolkit itself</b> — what each feature does, how a user drives it, what it "
          "computes, and the software design behind it. The statistical findings for the "
          "demonstration dataset are reported separately in the companion <i>Analysis "
          "Documentation</i>."),
        H2("1.2  Problem the system solves"),
        PROB("Producing a statistical report normally means repeating the same manual steps — "
             "loading data, running tests, making charts, writing up findings — for every "
             "dataset. Smart Analysis Reporter packages those steps into reusable tools and a "
             "report engine, so analysis and reporting become repeatable and fast."),
        H2("1.3  Intended audience"),
        BUL(["Evaluators assessing the toolkit's capabilities and engineering.",
             "Future maintainers extending the analysis modules or adding features.",
             "Analysts who want to understand exactly what each tool computes."]),
        H2("1.4  Document map"),
        P("Section 2 gives an abstract; Section 3 covers the system architecture and data "
          "flow; Section 4 the technology stack; Section 5 the codebase and module design; "
          "Section 6 is a granular catalogue of every feature; Section 7 details the report "
          "generation subsystem; Section 8 introduces the bundled demonstration dataset; and "
          "Section 9 concludes with future work."),
    ))

    # 2 ── Abstract
    sections.append(SEC(
        "2. Abstract",
        P("Smart Analysis Reporter is a layered, modular Python application with a Streamlit "
          "multipage front end and a dedicated report engine. It provides nine interactive "
          "analysis tools — data profiling, descriptive statistics, frequency tables, Q-Q "
          "normality plots, correlation, hypothesis testing, regression, time series — plus a "
          "report builder that turns any analysis into an editable, exportable document "
          "(PDF, Word, HTML). Analysis logic lives in pure-Python modules independent of the "
          "UI, which keeps the system testable, reusable and easy to extend."),
    ))

    # 3 ── System architecture & design
    sections.append(SEC(
        "3. System Architecture & Design",
        H2("3.1  Architectural overview"),
        P("The application is organised in four layers with a strict separation of concerns: "
          "a <b>data layer</b> that loads and caches the dataset; an <b>analysis layer</b> of "
          "pure-Python modules that implement the statistics; a <b>presentation layer</b> of "
          "Streamlit pages; and a <b>reporting layer</b> that renders charts and exports "
          "documents. Pages never contain statistical logic — they call modules — which means "
          "the same functions power both the interactive app and the generated documentation."),
        FIG(architecture_png()),
        CAP("Figure 3.1 — System architecture and data flow."),
        H2("3.2  Runtime data flow"),
        BUL(["On first load, the dataset is read once and stored in Streamlit "
             "<i>session state</i> (key <b>df</b>), shared across every page.",
             "Selecting a page runs its script: it reads <b>df</b>, calls the relevant analysis "
             "module, and renders interactive Plotly charts and tables.",
             "An “➕ Add to report” control renders a themed Matplotlib image of the current "
             "output, writes an auto-generated inference, and appends an item to the report "
             "state (key <b>report</b>).",
             "The Report Builder reads that state for editing; the export engine serialises it "
             "to PDF / DOCX / HTML."]),
        H2("3.3  State management"),
        P("Two objects in session state hold all cross-page state: <b>df</b> (the active "
          "dataset) and <b>report</b> (a dictionary of a <i>cover</i> plus an ordered list of "
          "<i>items</i>). Because Streamlit re-runs a page top-to-bottom on every interaction, "
          "test results are also stashed in session state so they survive re-runs and remain "
          "available to add to the report."),
        H2("3.4  Design principles"),
        BUL(["<b>Modularity</b> — statistics live in importable modules, not in page scripts.",
             "<b>Reusability</b> — chart renderers and inference builders are shared by the app, "
             "the report engine and these documents.",
             "<b>Portability</b> — the demonstration dataset is cached in-repo, so the deployed "
             "app needs no external credentials.",
             "<b>Separation of UI and logic</b> — enabling headless testing of every module."]),
    ))

    # 4 ── Technology stack
    stack = pd.DataFrame({
        "Technology": ["Streamlit", "pandas, NumPy", "SciPy, statsmodels", "Plotly",
                       "Matplotlib", "ReportLab", "python-docx", "kagglehub"],
        "Role in the toolkit": [
            "Multipage interactive web UI, widgets and session state",
            "Data structures and vectorised computation",
            "Statistical tests, distributions and OLS / ARIMA models",
            "Interactive on-screen charts",
            "Themed static charts embedded in exported documents",
            "PDF report and documentation generation",
            "Editable Word (.docx) generation",
            "Fetching / caching the demonstration dataset"],
    }, index=["UI", "Computation", "Statistics", "Interactive charts",
              "Report charts", "PDF export", "Word export", "Data source"])
    stack.index.name = "Layer"
    sections.append(SEC(
        "4. Technology Stack",
        P("The toolkit is built entirely in Python. Each dependency maps to a clear role:"),
        TBL(stack),
    ))

    # 5 ── Codebase & module design
    modmap = pd.DataFrame({"Responsibility": [
        "Home page — dataset loading, overview, one-click report trigger",
        "Theme (white/orange), page header, KPI cards, Plotly template",
        "Cached loader for the default demonstration dataset",
        "File upload parsing and Metric/Ordinal/Nominal column classification",
        "Dataset summary and per-column profiling",
        "Summary statistics (centre, spread, shape)",
        "Categorical and grouped-metric frequency tables",
        "Q-Q plotting positions and Shapiro-Wilk normality",
        "Pearson / Spearman / Kendall correlation matrices",
        "t-tests, ANOVA, chi-square, Z-tests and effect sizes",
        "Live test assumption checks (normality, variance, sample size)",
        "Simple and multiple OLS regression",
        "Time-series preparation and ARIMA forecast",
        "Report state, themed chart rendering, PDF/DOCX/HTML export",
        "One-click full-analysis report builder",
        "Analysis & System documentation generators",
    ]}, index=[
        "app.py", "ui.py", "datasets.py", "data_loader.py", "profiling.py", "descriptive.py",
        "frequency.py", "normality.py", "correlation.py", "tests.py", "assumptions.py",
        "regression.py", "timeseries.py", "report.py", "autoreport.py", "docgen.py"])
    modmap.index.name = "Module"
    sections.append(SEC(
        "5. Codebase & Module Design",
        P("The application is split into a thin set of Streamlit page scripts (in <b>pages/</b>) "
          "and a library of analysis modules (in <b>modules/</b>). Each module owns one "
          "responsibility and exposes plain functions that take a pandas object and return "
          "numbers, data frames or chart images — never UI. This table maps every module to its "
          "responsibility:"),
        TBL(modmap),
    ))

    # 6 ── Feature catalogue (granular, per tool)
    feat = [P("This section documents each tool in the toolkit. For every feature it states the "
              "purpose, the controls the user drives, the methods and statistics computed, the "
              "outputs produced, and the module that implements it.")]

    feat += _feature(
        "6.1  Home & Data Loading",
        "Entry point of the app; makes a dataset available to every other tool and offers a "
        "one-click report.",
        "Upload a CSV/Excel file, or use the bundled default dataset; reset to default.",
        "Files are parsed with pandas (read_csv / read_excel); columns are classified as "
        "Metric, Ordinal (low-cardinality integers) or Nominal by a dtype + cardinality "
        "heuristic. The dataset is held in session state and cached.",
        "KPI cards (rows, columns, missing, duplicates), a data preview, a column-type table, "
        "and an “Auto-generate report” button.",
        "modules/datasets.py, modules/data_loader.py (app.py)")

    feat += _feature(
        "6.2  Data Profiler",
        "A one-glance data-quality and structure report.",
        "None — operates on the loaded dataset.",
        "Computes row/column counts, total missing values and duplicate rows; per column it "
        "reports dtype, missing percentage, unique count and measurement level.",
        "KPI cards plus a merged profile table (dtype + measurement level) and a preview.",
        "modules/profiling.py, modules/data_loader.py")

    feat += _feature(
        "6.3  Descriptive Statistics",
        "Summarise the centre, spread and shape of each metric variable.",
        "A variable selector to inspect one column in detail.",
        "Mean, median, mode, variance, standard deviation, min/max/range, quartiles and IQR, "
        "plus skewness and kurtosis (pandas / SciPy). A fitted normal curve is overlaid on the "
        "histogram.",
        "A summary table for all metric variables, KPI cards, a histogram with normal overlay, "
        "and a box plot.",
        "modules/descriptive.py")
    if target:
        feat += [FIG(R.hist_png(df[target], title=f"Distribution of {target}"), 0.66),
                 CAP("Figure 6.1 — Example output of the Descriptive tool: histogram with "
                     "fitted normal curve.")]

    feat += _feature(
        "6.4  Frequency Tables",
        "Show how often each category or binned value occurs.",
        "A categorical variable, or a metric variable with a bin-count slider.",
        "Absolute, relative (%) and cumulative-(%) frequencies via value counts; metric "
        "variables are binned with equal-width intervals (pd.cut).",
        "A frequency table and an accompanying bar chart.",
        "modules/frequency.py")

    feat += _feature(
        "6.5  Q-Q Plot & Normality",
        "Test whether a variable is normally distributed — a key assumption for parametric tests.",
        "A metric variable.",
        "Blom plotting positions are compared against theoretical normal quantiles "
        "(SciPy norm.ppf) with a Q1–Q3 reference line; a Shapiro-Wilk test and skewness / "
        "kurtosis quantify the departure from normality.",
        "A Q-Q plot, a histogram with normal curve, the W statistic and p-value, KPI cards and "
        "a plain-language verdict.",
        "modules/normality.py")
    if target:
        qq = qq_data(df[target])
        if qq is not None:
            feat += [FIG(R.qq_png(qq["theoretical"], qq["sample"], qq["ref_x"], qq["ref_y"]), 0.66),
                     CAP("Figure 6.2 — Example output of the Q-Q tool: sample quantiles against "
                         "the normal reference line.")]

    feat += _feature(
        "6.6  Correlation",
        "Measure the linear association between every pair of metric variables.",
        "A correlation method — Pearson, Spearman or Kendall.",
        "The full correlation matrix is computed and rendered as a colour heatmap; the "
        "strongest pair is summarised automatically.",
        "A correlation-matrix table, a heatmap, and an auto-written interpretation.",
        "modules/correlation.py")

    feat += _feature(
        "6.7  Hypothesis Testing",
        "Decide whether observed differences or associations are statistically significant "
        "(α = 0.05).",
        "Choice of test (one-sample t, Welch independent t, paired t, ANOVA, chi-square, "
        "one- and two-sample Z); the relevant variables, hypothesised mean μ₀ and σ.",
        "Tests use SciPy; the independent t-test applies Welch's correction; effect sizes "
        "(Cohen's d, Cramér's V) and 95% confidence intervals are reported. A <b>live "
        "assumption engine</b> checks normality (Shapiro-Wilk), equal variances (Levene), "
        "sample-size adequacy (CLT) and expected-cell counts, shown as pass / review flags "
        "before the test is run.",
        "Test statistic, p-value, effect size, contingency table (chi-square) and a "
        "plain-language verdict; results persist across re-runs and can be added to the report.",
        "modules/tests.py, modules/assumptions.py")

    feat += _feature(
        "6.8  Regression",
        "Quantify and model how predictors relate to an outcome.",
        "Simple linear: an X (predictor) and Y (outcome). Multiple OLS: a dependent variable "
        "and several independent variables.",
        "Ordinary-least-squares fitting via statsmodels; the simple model reports intercept, "
        "slope, R² and the slope's p-value, and draws the fitted line; multiple OLS returns the "
        "full regression summary.",
        "The fitted equation, KPI cards, an R² indicator, a scatter-plus-fit chart, and the OLS "
        "summary table.",
        "modules/regression.py")
    if target and predictor:
        rr = simple_linear_regression(df, predictor, target)
        if rr:
            feat += [FIG(R.regression_png(rr["x"], rr["y"], rr["x_line"], rr["y_line"],
                                          xlabel=predictor, ylabel=target,
                                          title=f"{target} vs {predictor}"), 0.66),
                     CAP("Figure 6.3 — Example output of the Regression tool: data with the "
                         "fitted least-squares line.")]

    feat += _feature(
        "6.9  Time Series",
        "Reveal a trend and produce a short forecast for date-bearing data.",
        "A date column and a value column.",
        "Dates are parsed, sorted and indexed; an ARIMA(1,1,1) model produces a 30-step "
        "forecast (statsmodels). The tool gracefully reports when the dataset has no date "
        "column.",
        "A trend line chart and a forecast table.",
        "modules/timeseries.py")

    sections.append(SEC("6. Feature Catalogue", *feat))

    # 7 ── Report generation subsystem
    sections.append(SEC(
        "7. Report Generation Subsystem",
        P("The reporting layer is what makes the toolkit a <i>reporter</i>. It is built around a "
          "single state object and a set of renderers."),
        H2("7.1  Report state model"),
        P("The report is a dictionary held in session state with two parts: a <b>cover</b> "
          "(title, subtitle, author, ID, date and an editable executive summary) and an ordered "
          "list of <b>items</b>. Each item carries an id, a title, an editable inference, an "
          "optional chart image (PNG) and an optional table (data frame)."),
        H2("7.2  Add-to-report (incremental capture)"),
        P("Every analysis page exposes an “➕ Add to report” control. When used, the toolkit "
          "renders a themed Matplotlib image of the current chart, generates a one-sentence "
          "inference from the underlying statistics, and appends an item (de-duplicated by "
          "title). The user thus assembles a report while exploring."),
        H2("7.3  Auto-report (one-click)"),
        P("The auto-report builder chains the analysis modules to produce a complete report in "
          "one click — dataset overview, descriptive statistics, target distribution, spread "
          "and outliers, correlation, best-predictor regression, a categorical breakdown, a "
          "group comparison test and a normality check — each with an auto-written inference."),
        H2("7.4  Report Builder page"),
        BUL(["Edit the cover — title, subtitle, author, ID, date and executive summary.",
             "Edit each section's inference in place; reorder (↑/↓) or remove (✕) sections.",
             "Preview the orange cover banner live before exporting."]),
        H2("7.5  Export engine"),
        P("A single report model is serialised three ways: a paginated <b>PDF</b> (ReportLab) "
          "with an orange cover and embedded charts/tables; an editable <b>Word</b> document "
          "(python-docx); and a self-contained <b>HTML</b> file with base64-embedded images. "
          "The same renderers power these companion documentation files.",),
        H2("7.6  Backing modules"),
        P("modules/report.py (state, chart rendering, exporters), modules/autoreport.py "
          "(one-click builder), modules/docgen.py (these documents)."),
    ))

    # 8 ── Demonstration dataset (brief — analysis is separate)
    cols = pd.DataFrame({
        "Type": ["Categorical" if c in categorical else "Metric" for c in df.columns],
        "Example": [str(df[c].iloc[0]) for c in df.columns],
        "Unique": [int(df[c].nunique()) for c in df.columns],
    }, index=list(df.columns))
    cols.index.name = "Column"
    sections.append(SEC(
        "8. Demonstration Dataset",
        P(f"So the toolkit works out of the box, it bundles the <b>{name}</b> dataset "
          f"({len(df):,} rows × {df.shape[1]} columns, target <b>{target}</b>). It is used here "
          f"only to illustrate the features — the full statistical study of this data is "
          f"presented in the companion <i>Analysis Documentation</i>. The schema is:"),
        TBL(cols),
    ))

    # 9 ── Project conclusion
    sections.append(SEC(
        "9. Project Conclusion",
        P("Smart Analysis Reporter packages a complete analytics workflow — from raw data, "
          "through a suite of interactive statistical tools, to a polished and editable report "
          "— behind a clean white-and-orange interface. Its layered design keeps statistics, "
          "UI and reporting independent, so new analyses, datasets or export formats can be "
          "added with minimal change, and every module can be tested headlessly."),
        H2("9.1  Future work"),
        BUL(["Multiple-regression and classification tools with model diagnostics.",
             "Automated outlier and data-quality reporting.",
             "Saved report templates and user accounts.",
             "Additional export themes and chart types."]),
    ))

    meta = {"title": "System Documentation", "doc_kind": "Toolkit & System Design Report",
            "subtitle": name}
    return meta, sections


# ═══════════════════ DEPLOYMENT GUIDE MODEL ═══════════════════

REPO_URL = "https://github.com/TejalMenezes/analysis-toolkit-v1"


def build_deploy_model(repo_url=REPO_URL):
    repo_git = repo_url + ".git"
    sections = []

    sections.append(SEC(
        "1. Overview",
        P("This guide explains how to publish <b>Smart Analysis Reporter</b> as a live, public "
          "web application. It is organised as two independent process flows:"),
        BUL(["<b>Flow A — Push the project to GitHub.</b>",
             "<b>Flow B — Deploy from GitHub to a public link</b> (Streamlit Community Cloud)."]),
        P(f"Target repository: <b>{repo_url}</b> (branch <b>main</b>)."),
        H2("1.1  Pre-flight checklist"),
        P("These are already prepared in the project:"),
        BUL(["requirements.txt lists every dependency.",
             ".gitignore excludes the virtual environment, caches, logs and local settings.",
             "The default dataset is committed at data/student_dataset.csv, so the deployed app "
             "needs no Kaggle login.",
             "The application entry point is app.py.",
             "The Git repository is initialised and committed locally."]),
    ))

    sections.append(SEC(
        "2. Flow A — Push to GitHub",
        P("Run these commands from the project root. They point the local repository at GitHub "
          "and upload the code."),
        CODE("# 1. point the local repo at your GitHub repo (once)\n"
             f"git remote add origin {repo_git}\n\n"
             "# 2. make sure you are on main\n"
             "git branch -M main\n\n"
             "# 3. stage and commit any pending changes\n"
             "git add .\n"
             'git commit -m "Smart Analysis Reporter"\n\n'
             "# 4. push\n"
             "git push -u origin main"),
        H2("2.1  If the remote already exists"),
        P("If Git reports “remote origin already exists”, update its URL instead:"),
        CODE(f"git remote set-url origin {repo_git}\n"
             "git push -u origin main"),
        H2("2.2  Authentication"),
        P("GitHub no longer accepts your account password on the command line — use a "
          "<b>Personal Access Token (PAT)</b>:"),
        BUL(["In GitHub, go to Settings → Developer settings → Personal access tokens → "
             "Tokens (classic).",
             "Generate a new token with the <b>repo</b> scope and copy it.",
             "When <i>git push</i> asks for a password, paste the token.",
             "Alternatively, use GitHub Desktop or the GitHub CLI (gh auth login), which "
             "handle authentication for you."]),
        P("<b>Note on multiple accounts.</b> If the machine has cached credentials for a "
          "different GitHub account, the push will be rejected with a 403 error. Clear the "
          "stored credential (Windows Credential Manager → remove the github.com entry) or "
          "push using a token for the correct account."),
        H2("2.3  Verify"),
        P(f"Refresh {repo_url} in your browser — all the project files should now appear."),
    ))

    sections.append(SEC(
        "3. Flow B — Deploy to a Public Link",
        P("Once the repository is on GitHub, Streamlit Community Cloud builds and hosts the app "
          "for free — there is no server to manage."),
        H2("3.1  Steps"),
        BUL(["Go to https://share.streamlit.io and sign in with GitHub (authorise repo access).",
             "Click “Create app” → “Deploy a public app from GitHub.”",
             "Repository: <b>TejalMenezes/analysis-toolkit-v1</b>; Branch: <b>main</b>; "
             "Main file path: <b>app.py</b>.",
             "Click “Deploy.”",
             "Wait for the first build — it installs requirements.txt (typically 2–5 minutes; "
             "the log streams live).",
             "When the build finishes you receive your public URL, for example "
             "https://analysis-toolkit-v1.streamlit.app — that is the link to share."]),
        H2("3.2  After deployment"),
        BUL(["Every push to main automatically redeploys the app — no extra steps.",
             "Manage, reboot or view logs from the dashboard at https://share.streamlit.io.",
             "If a build fails, open the log — it is almost always a missing line in "
             "requirements.txt."]),
    ))

    sections.append(SEC(
        "4. Updating the Live App",
        P("To publish a change after the app is live, commit and push — Streamlit Cloud detects "
          "the push and rebuilds automatically."),
        CODE("git add .\n"
             'git commit -m "describe your change"\n'
             "git push"),
    ))

    trouble = pd.DataFrame({"Resolution": [
        "Run: git remote set-url origin <url>, then push.",
        "Run: git pull --rebase origin main, then git push.",
        "Use a Personal Access Token as the password (see Section 2.2).",
        "Add the missing package to requirements.txt, commit and push.",
        "Confirm data/student_dataset.csv is committed (it is, by default).",
    ]}, index=[
        "remote origin already exists",
        "Push rejected (non-fast-forward)",
        "Authentication fails on push",
        "Cloud build fails on an import",
        "App loads but reports no dataset"])
    trouble.index.name = "Symptom"
    sections.append(SEC(
        "5. Troubleshooting",
        P("Common issues and how to resolve them:"),
        TBL(trouble),
    ))

    sections.append(SEC(
        "6. Quick Temporary Link",
        P("If you only need a public link for a short demo from your own machine (it lives only "
          "while the command runs), expose the local app with a tunnel:"),
        CODE("# terminal 1\n"
             "streamlit run app.py\n\n"
             "# terminal 2\n"
             "pip install pyngrok\n"
             'python -c "from pyngrok import ngrok; print(ngrok.connect(8501))"'),
        P("For a permanent link, use Flow B instead."),
    ))

    meta = {"title": "Deployment Guide", "doc_kind": "Deployment Process Flow",
            "subtitle": "Smart Analysis Reporter"}
    return meta, sections


# ═══════════════════ PDF RENDERER ═══════════════════

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor, white
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT
from reportlab.platypus import (BaseDocTemplate, PageTemplate, Frame, Paragraph,
                                Spacer, Table, TableStyle, PageBreak, NextPageTemplate,
                                ListFlowable, ListItem)
from reportlab.platypus.tableofcontents import TableOfContents

_ss = getSampleStyleSheet()
_PS = {
    "h1": ParagraphStyle("SecH1", parent=_ss["Heading1"], textColor=HexColor(ORANGE_DARK),
                         fontSize=17, leading=21, spaceBefore=2, spaceAfter=10),
    "h2": ParagraphStyle("SecH2", parent=_ss["Heading2"], textColor=HexColor(INK),
                         fontSize=12.5, leading=16, spaceBefore=12, spaceAfter=4),
    "body": ParagraphStyle("Body", parent=_ss["BodyText"], fontSize=10.5, leading=15.5,
                           textColor=HexColor("#2A2A2A"), spaceAfter=8, alignment=TA_LEFT),
    "bullet": ParagraphStyle("Bullet", parent=_ss["BodyText"], fontSize=10.5, leading=15,
                             textColor=HexColor("#2A2A2A")),
    "cap": ParagraphStyle("Cap", parent=_ss["BodyText"], fontSize=8.5, leading=11,
                          textColor=HexColor(MUTED), spaceBefore=3, spaceAfter=12),
    "tocHead": ParagraphStyle("TocHead", parent=_ss["Heading1"], textColor=HexColor(INK),
                              fontSize=16, spaceAfter=14),
    "pbLabel": ParagraphStyle("PbLabel", parent=_ss["BodyText"], fontSize=8.5,
                              textColor=HexColor(ORANGE_DARK), spaceAfter=2),
    "pbText": ParagraphStyle("PbText", parent=_ss["BodyText"], fontSize=10.5, leading=15,
                             textColor=HexColor(INK)),
}
_AVAIL_W = A4[0] - 4 * cm


def _pb_box(text):
    inner = [[Paragraph("PROBLEM STATEMENT", _PS["pbLabel"])], [Paragraph(text, _PS["pbText"])]]
    t = Table(inner, colWidths=[_AVAIL_W])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor(ORANGE_SOFT)),
        ("LINEBEFORE", (0, 0), (0, -1), 3, HexColor(ORANGE)),
        ("LEFTPADDING", (0, 0), (-1, -1), 12), ("RIGHTPADDING", (0, 0), (-1, -1), 12),
        ("TOPPADDING", (0, 0), (0, 0), 9), ("BOTTOMPADDING", (-1, -1), (-1, -1), 10),
        ("TOPPADDING", (0, 1), (0, 1), 0),
    ]))
    return t


def _code_box(text):
    from reportlab.platypus import Preformatted
    code_style = ParagraphStyle("Code", fontName="Courier", fontSize=8.5, leading=11.5,
                                textColor=HexColor("#1F2430"))
    t = Table([[Preformatted(text, code_style)]], colWidths=[_AVAIL_W])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), HexColor("#F4F4F6")),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor("#E2E2E8")),
        ("LEFTPADDING", (0, 0), (-1, -1), 10), ("RIGHTPADDING", (0, 0), (-1, -1), 10),
        ("TOPPADDING", (0, 0), (-1, -1), 8), ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    return t


def _pdf_block(b):
    kind = b[0]
    if kind == "h2":      return _PS_para(b[1], "h2")
    if kind == "para":    return _PS_para(b[1], "body")
    if kind == "caption": return _PS_para(b[1], "cap")
    if kind == "problem": return _pb_box(b[1])
    if kind == "code":    return _code_box(b[1])
    if kind == "figure":  return R._rl_image(b[1], _AVAIL_W * b[2])
    if kind == "table":   return R._rl_table(b[1], _AVAIL_W)
    if kind == "bullets":
        return ListFlowable([ListItem(Paragraph(t, _PS["bullet"]), bulletColor=HexColor(ORANGE),
                                      value="•") for t in b[1]],
                            bulletType="bullet", start="•", leftIndent=14)
    return Spacer(1, 1)


def _PS_para(text, style):
    return Paragraph(text, _PS[style])


class _DocTemplate(BaseDocTemplate):
    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph) and flowable.style.name == "SecH1":
            self.notify("TOCEntry", (0, flowable.getPlainText(), self.page))


def render_pdf(meta, sections):
    buf = io.BytesIO()
    W, H = A4

    def draw_cover(canvas, doc):
        canvas.saveState()
        canvas.setFillColor(HexColor(ORANGE)); canvas.rect(0, H - 10.5 * cm, W, 10.5 * cm, fill=1, stroke=0)
        canvas.setFillColor(HexColor(ORANGE_DARK)); canvas.rect(0, H - 10.9 * cm, W, 0.4 * cm, fill=1, stroke=0)
        canvas.setFillColor(white)
        canvas.setFont("Helvetica-Bold", 12); canvas.drawString(2 * cm, H - 2.3 * cm, APP_NAME.upper())
        canvas.setFont("Helvetica", 10); canvas.drawString(2 * cm, H - 2.9 * cm, meta["doc_kind"])
        canvas.setFont("Helvetica-Bold", 30); canvas.drawString(2 * cm, H - 5.2 * cm, meta["title"])
        canvas.setFont("Helvetica", 14); canvas.drawString(2 * cm, H - 6.3 * cm, meta["subtitle"])
        canvas.setFont("Helvetica", 11)
        canvas.drawString(2 * cm, H - 8.6 * cm, f"Prepared by {AUTHOR_NAME}")
        canvas.drawString(2 * cm, H - 9.2 * cm, f"Student ID: {AUTHOR_ID}")
        canvas.drawString(2 * cm, H - 9.8 * cm, f"Date: {date.today().isoformat()}")
        canvas.setFillColor(HexColor(MUTED)); canvas.setFont("Helvetica", 8)
        canvas.drawString(2 * cm, 1.4 * cm, f"{APP_NAME} — Tools & Methods of Data Analysis")
        canvas.restoreState()

    def draw_body(canvas, doc):
        canvas.saveState()
        canvas.setStrokeColor(HexColor(ORANGE_LINE)); canvas.setLineWidth(1)
        canvas.line(2 * cm, H - 1.55 * cm, W - 2 * cm, H - 1.55 * cm)
        canvas.setFillColor(HexColor(MUTED)); canvas.setFont("Helvetica", 8)
        canvas.drawString(2 * cm, H - 1.4 * cm, meta["title"])
        canvas.drawRightString(W - 2 * cm, H - 1.4 * cm, APP_NAME)
        canvas.line(2 * cm, 1.5 * cm, W - 2 * cm, 1.5 * cm)
        canvas.drawString(2 * cm, 1.1 * cm, f"{AUTHOR_NAME} · ID {AUTHOR_ID}")
        canvas.drawRightString(W - 2 * cm, 1.1 * cm, f"Page {doc.page}")
        canvas.restoreState()

    doc = _DocTemplate(buf, pagesize=A4, title=meta["title"], author=AUTHOR_NAME,
                       leftMargin=2 * cm, rightMargin=2 * cm)
    doc.addPageTemplates([
        PageTemplate(id="cover", frames=[Frame(2 * cm, 2 * cm, W - 4 * cm, H - 4 * cm)], onPage=draw_cover),
        PageTemplate(id="body", frames=[Frame(2 * cm, 1.8 * cm, W - 4 * cm, H - 3.8 * cm)], onPage=draw_body),
    ])

    toc = TableOfContents()
    toc.levelStyles = [ParagraphStyle("toc0", fontSize=11, leading=20, textColor=HexColor(INK))]

    story = [Spacer(1, 1), NextPageTemplate("body"), PageBreak(),
             Paragraph("Table of Contents", _PS["tocHead"]), toc, PageBreak()]

    for sec in sections:
        story.append(Paragraph(sec["title"], _PS["h1"]))
        for b in sec["blocks"]:
            story.append(_pdf_block(b))
        story.append(PageBreak())

    doc.multiBuild(story)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════ DOCX RENDERER ═══════════════════

def _docx_runs(paragraph, text):
    bold = False
    for tok in re.split(r"(<b>|</b>)", text):
        if tok == "<b>":
            bold = True
        elif tok == "</b>":
            bold = False
        elif tok:
            run = paragraph.add_run(tok)
            run.bold = bold


def _shade(cell_or_para_pr, hex_color):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    cell_or_para_pr.append(shd)


def render_docx(meta, sections):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    O = RGBColor(0xF5, 0x7C, 0x00)
    OD = RGBColor(0xE6, 0x51, 0x00)
    GREY = RGBColor(0x7A, 0x82, 0x90)

    doc = Document()
    # base styling
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)

    sec0 = doc.sections[0]
    sec0.top_margin = Cm(2.2); sec0.bottom_margin = Cm(2)
    sec0.left_margin = Cm(2.2); sec0.right_margin = Cm(2.2)

    # footer with page number
    footer = sec0.footer
    fp = footer.paragraphs[0]
    fp.text = f"{AUTHOR_NAME} · ID {AUTHOR_ID}        "
    fp.runs[0].font.size = Pt(8); fp.runs[0].font.color.rgb = GREY
    run = fp.add_run("Page ")
    run.font.size = Pt(8); run.font.color.rgb = GREY
    # PAGE field
    fld1 = OxmlElement("w:fldSimple"); fld1.set(qn("w:instr"), "PAGE")
    fp._p.append(fld1)

    # header with doc title
    hp = sec0.header.paragraphs[0]
    hp.text = meta["title"]
    hp.runs[0].font.size = Pt(8); hp.runs[0].font.color.rgb = GREY

    # ── cover ──
    band = doc.add_paragraph()
    _shade(band._p.get_or_add_pPr(), "F57C00")
    band.paragraph_format.space_after = Pt(2)
    r = band.add_run(APP_NAME.upper()); r.bold = True; r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.size = Pt(12)
    k = doc.add_paragraph(); _shade(k._p.get_or_add_pPr(), "F57C00")
    rk = k.add_run(meta["doc_kind"]); rk.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rk.font.size = Pt(10)
    tt = doc.add_paragraph(); _shade(tt._p.get_or_add_pPr(), "F57C00")
    rt = tt.add_run(meta["title"]); rt.bold = True; rt.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rt.font.size = Pt(28)
    sb = doc.add_paragraph(); _shade(sb._p.get_or_add_pPr(), "F57C00")
    rs = sb.add_run(meta["subtitle"]); rs.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); rs.font.size = Pt(14)
    sb.paragraph_format.space_after = Pt(16)

    for line in [f"Prepared by {AUTHOR_NAME}", f"Student ID: {AUTHOR_ID}", f"Date: {date.today().isoformat()}"]:
        p = doc.add_paragraph(); rr = p.add_run(line); rr.font.size = Pt(11)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── table of contents (Word field; updates on open) ──
    th = doc.add_paragraph(); rth = th.add_run("Table of Contents"); rth.bold = True; rth.font.size = Pt(16)
    tocp = doc.add_paragraph()
    fb = OxmlElement("w:r"); fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin"); fb.append(fc); tocp._p.append(fb)
    it = OxmlElement("w:r"); itx = OxmlElement("w:instrText"); itx.set(qn("xml:space"), "preserve")
    itx.text = 'TOC \\o "1-1" \\h \\z \\u'; it.append(itx); tocp._p.append(it)
    fs = OxmlElement("w:r"); fsc = OxmlElement("w:fldChar"); fsc.set(qn("w:fldCharType"), "separate"); fs.append(fsc); tocp._p.append(fs)
    pl = OxmlElement("w:r"); pt = OxmlElement("w:t"); pt.text = "Right-click and ‘Update Field’ to build the table of contents."
    pl.append(pt); tocp._p.append(pl)
    fe = OxmlElement("w:r"); fec = OxmlElement("w:fldChar"); fec.set(qn("w:fldCharType"), "end"); fe.append(fec); tocp._p.append(fe)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ── sections ──
    for si, sec in enumerate(sections):
        if si > 0:
            doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
        hh = doc.add_heading(level=1)
        rhh = hh.add_run(sec["title"]); rhh.font.color.rgb = OD
        for b in sec["blocks"]:
            kind = b[0]
            if kind == "h2":
                h = doc.add_heading(level=2); rh = h.add_run(b[1]); rh.font.color.rgb = RGBColor(0x1F, 0x24, 0x30)
            elif kind == "para":
                _docx_runs(doc.add_paragraph(), b[1])
            elif kind == "caption":
                p = doc.add_paragraph(); rr = p.add_run(b[1]); rr.italic = True
                rr.font.size = Pt(8.5); rr.font.color.rgb = GREY
            elif kind == "code":
                cp = doc.add_paragraph(); _shade(cp._p.get_or_add_pPr(), "F4F4F6")
                cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(4)
                lines = b[1].split("\n")
                for li, line in enumerate(lines):
                    if li:
                        cp.add_run().add_break()
                    rc = cp.add_run(line)
                    rc.font.name = "Consolas"; rc.font.size = Pt(8.5)
                    rc.font.color.rgb = RGBColor(0x1F, 0x24, 0x30)
            elif kind == "bullets":
                for item in b[1]:
                    _docx_runs(doc.add_paragraph(style="List Bullet"), item)
            elif kind == "problem":
                lab = doc.add_paragraph(); _shade(lab._p.get_or_add_pPr(), "FFF6EE")
                rl = lab.add_run("PROBLEM STATEMENT"); rl.bold = True; rl.font.size = Pt(8.5); rl.font.color.rgb = OD
                lab.paragraph_format.space_after = Pt(0)
                body = doc.add_paragraph(); _shade(body._p.get_or_add_pPr(), "FFF6EE")
                _docx_runs(body, b[1])
            elif kind == "figure":
                doc.add_picture(io.BytesIO(b[1]), width=Inches(6.2 * b[2]))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif kind == "table":
                _docx_table(doc, b[1], O)

    buf = io.BytesIO()
    doc.save(buf); buf.seek(0)
    return buf.getvalue()


def _docx_table(doc, df, header_rgb):
    from docx.shared import Pt, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    t = doc.add_table(rows=1, cols=len(df.columns) + 1)
    t.style = "Light Grid Accent 6"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    hdr[0].text = str(df.index.name or "")
    for j, c in enumerate(df.columns):
        hdr[j + 1].text = str(c)
    for cell in hdr:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for idx, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(idx)
        for j, c in enumerate(df.columns):
            v = row[c]
            cells[j + 1].text = (f"{v:.3f}" if isinstance(v, (int, float, np.floating))
                                 and not isinstance(v, bool) else str(v))
        for cell in cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(8.5)


# ═══════════════════ public builders ═══════════════════

def build_analysis_doc(df, name):
    meta, sections = build_analysis_model(df, name)
    return render_pdf(meta, sections)


def build_analysis_docx(df, name):
    meta, sections = build_analysis_model(df, name)
    return render_docx(meta, sections)


def build_system_doc(df, name):
    meta, sections = build_system_model(df, name)
    return render_pdf(meta, sections)


def build_system_docx(df, name):
    meta, sections = build_system_model(df, name)
    return render_docx(meta, sections)


def build_deploy_doc():
    meta, sections = build_deploy_model()
    return render_pdf(meta, sections)


def build_deploy_docx():
    meta, sections = build_deploy_model()
    return render_docx(meta, sections)
